import streamlit as st
import os
import time
import re
from io import BytesIO
import base64
import difflib
from collections import Counter

# Document processing
import PyPDF2
import docx

# Summarization engines
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words

# Transformers for deep summarization
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import torch

# Translation
from googletrans import Translator

# Wikipedia
import wikipedia

# NLP utilities
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Export utilities
from fpdf import FPDF
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

# Set page config
st.set_page_config(
    page_title="Universal AI Summarizer v4.1 Turbo",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 1rem 0;
    }
    .stButton>button {
        width: 100%;
        background-color: #667eea;
        color: white;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        font-weight: bold;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #667eea;
    }
    .footer {
        text-align: center;
        padding: 2rem 0 1rem 0;
        color: #666;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

# Cache models to avoid reloading
@st.cache_resource
def load_summarization_model(model_name="sshleifer/distilbart-cnn-12-6"):
    """Load and cache the transformer summarization model"""
    try:
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, device=-1)
        return summarizer
    except Exception as e:
        st.warning(f"Could not load {model_name}: {e}")
        return None

@st.cache_resource
def load_translator():
    """Load and cache the translator"""
    return Translator()

# Initialize session state
if 'summarization_history' not in st.session_state:
    st.session_state.summarization_history = []

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ""

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    try:
        text = file.read().decode('utf-8')
        return text
    except Exception as e:
        st.error(f"Error reading TXT: {e}")
        return ""

def preprocess_text(text):
    """Clean and preprocess text"""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep punctuation
    text = re.sub(r'[^\w\s.,!?;:\'-]', '', text)
    return text.strip()

def detect_language(text):
    """Detect language of text"""
    try:
        translator = load_translator()
        detection = translator.detect(text[:500])  # Detect from first 500 chars
        return detection.lang
    except:
        return 'en'

def translate_text(text, target_lang='en'):
    """Translate text to target language"""
    try:
        translator = load_translator()
        # Split into chunks if text is too long
        max_length = 4500
        if len(text) > max_length:
            chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            translated_chunks = []
            for chunk in chunks:
                result = translator.translate(chunk, dest=target_lang)
                translated_chunks.append(result.text)
            return ' '.join(translated_chunks)
        else:
            result = translator.translate(text, dest=target_lang)
            return result.text
    except Exception as e:
        st.warning(f"Translation failed: {e}")
        return text

def turbo_summarize(text, num_sentences=5):
    """Fast summarization using Sumy LSA"""
    try:
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        stemmer = Stemmer("english")
        summarizer = LsaSummarizer(stemmer)
        summarizer.stop_words = get_stop_words("english")
        
        summary_sentences = summarizer(parser.document, num_sentences)
        summary = " ".join([str(sentence) for sentence in summary_sentences])
        return summary
    except Exception as e:
        st.error(f"Turbo summarization error: {e}")
        return text[:500] + "..."

def deep_summarize(text, max_length=150, min_length=50):
    """Deep summarization using transformer models"""
    try:
        summarizer = load_summarization_model()
        if summarizer is None:
            return turbo_summarize(text)
        
        # Chunk text if too long (max ~1024 tokens for DistilBART)
        max_chunk_length = 1000
        words = text.split()
        
        if len(words) > max_chunk_length:
            chunks = [' '.join(words[i:i+max_chunk_length]) for i in range(0, len(words), max_chunk_length)]
            summaries = []
            
            for chunk in chunks:
                if len(chunk.split()) > 50:  # Only summarize substantial chunks
                    summary = summarizer(chunk, max_length=max_length, min_length=min_length, do_sample=False)
                    summaries.append(summary[0]['summary_text'])
            
            final_summary = ' '.join(summaries)
            
            # If combined summaries are still long, summarize again
            if len(final_summary.split()) > max_chunk_length:
                final_summary = summarizer(final_summary, max_length=max_length, min_length=min_length, do_sample=False)[0]['summary_text']
            
            return final_summary
        else:
            summary = summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
            return summary[0]['summary_text']
    except Exception as e:
        st.error(f"Deep summarization error: {e}")
        return turbo_summarize(text)

def analyze_tone(text):
    """Analyze the tone/sentiment of text"""
    # Simple keyword-based tone analysis
    positive_words = ['good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'positive', 'happy', 'success']
    negative_words = ['bad', 'terrible', 'awful', 'horrible', 'negative', 'sad', 'failure', 'poor', 'wrong']
    neutral_words = ['information', 'data', 'report', 'study', 'analysis', 'research', 'document']
    
    text_lower = text.lower()
    words = text_lower.split()
    
    pos_count = sum(1 for word in words if word in positive_words)
    neg_count = sum(1 for word in words if word in negative_words)
    neu_count = sum(1 for word in words if word in neutral_words)
    
    if pos_count > neg_count and pos_count > neu_count:
        return "Positive 😊"
    elif neg_count > pos_count and neg_count > neu_count:
        return "Negative 😟"
    elif neu_count > pos_count and neu_count > neg_count:
        return "Informative 📊"
    else:
        return "Neutral 😐"

def extract_keywords(text, top_n=10):
    """Extract top keywords using TF-IDF"""
    try:
        # Remove stopwords and short words
        words = text.lower().split()
        words = [w for w in words if len(w) > 4]
        
        if len(words) < top_n:
            return words
        
        # Use Counter for frequency
        word_freq = Counter(words)
        return [word for word, count in word_freq.most_common(top_n)]
    except Exception as e:
        return []

def fetch_wikipedia_context(query, sentences=3):
    """Fetch context from Wikipedia"""
    try:
        summary = wikipedia.summary(query, sentences=sentences, auto_suggest=True)
        return summary
    except wikipedia.exceptions.DisambiguationError as e:
        try:
            # Try first option
            summary = wikipedia.summary(e.options[0], sentences=sentences)
            return summary
        except:
            return None
    except:
        return None

def generate_qa_notes(text, summary):
    """Generate Q&A study notes from summary"""
    sentences = summary.split('.')
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    qa_pairs = []
    question_templates = [
        "What is the main point about {}?",
        "How does {} relate to the topic?",
        "Why is {} important?",
        "What can we learn from {}?",
        "Explain the significance of {}."
    ]
    
    for i, sentence in enumerate(sentences[:5]):  # Max 5 Q&A pairs
        # Extract key phrase (simple heuristic)
        words = sentence.split()
        if len(words) > 5:
            key_phrase = ' '.join(words[:5]) + "..."
        else:
            key_phrase = sentence
        
        question = question_templates[i % len(question_templates)].format("this point")
        answer = sentence + "."
        
        qa_pairs.append({"question": f"Q{i+1}: {question}", "answer": f"A{i+1}: {answer}"})
    
    return qa_pairs

def calculate_similarity(text1, text2):
    """Calculate cosine similarity between two texts"""
    try:
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([text1, text2])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return similarity * 100
    except:
        return 0.0

def generate_diff_html(text1, text2):
    """Generate HTML diff between two texts"""
    diff = difflib.ndiff(text1.split(), text2.split())
    html_diff = []
    
    for word in diff:
        if word.startswith('+ '):
            html_diff.append(f'<span style="background-color: #90EE90;">{word[2:]}</span>')
        elif word.startswith('- '):
            html_diff.append(f'<span style="background-color: #FFB6C6;">{word[2:]}</span>')
        elif word.startswith('  '):
            html_diff.append(word[2:])
    
    return ' '.join(html_diff)

def export_to_txt(content, filename="summary.txt"):
    """Export content to TXT"""
    return content.encode('utf-8')

def export_to_docx(content, filename="summary.docx"):
    """Export content to DOCX"""
    doc = DocxDocument()
    
    # Add title
    title = doc.add_heading('AI Summary Report', 0)
    title.alignment = 1  # Center
    
    # Add content
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_pdf(content, filename="summary.pdf"):
    """Export content to PDF"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "AI Summary Report", ln=True, align='C')
        pdf.ln(10)
        
        # Add content
        pdf.set_font("Arial", size=11)
        for line in content.split('\n'):
            if line.strip():
                pdf.multi_cell(0, 10, line)
                pdf.ln(2)
        
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

# Main App
def main():
    # Header
    st.markdown('<p class="main-header">🧠 Universal AI Summarizer v4.1 Turbo Edition</p>', unsafe_allow_html=True)
    st.markdown("---")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        
        mode = st.radio(
            "Summarization Mode",
            ["🚀 Turbo (Fast - LSA)", "🧠 Deep (Transformer)"],
            help="Turbo mode is faster but less accurate. Deep mode uses AI transformers."
        )
        
        turbo_mode = "Turbo" in mode
        
        if not turbo_mode:
            summary_length = st.slider("Summary Length", 50, 300, 150, help="Max words in summary")
        else:
            num_sentences = st.slider("Number of Sentences", 3, 10, 5, help="Number of sentences in summary")
        
        enable_translation = st.checkbox("🌍 Auto-translate to English", value=False)
        enable_wikipedia = st.checkbox("📚 Wikipedia Enrichment", value=False)
        
        if enable_wikipedia:
            wiki_query = st.text_input("Wikipedia Topic (optional)", placeholder="Enter topic...")
        
        st.markdown("---")
        st.header("📤 Export Options")
        export_format = st.multiselect("Select formats", ["TXT", "DOCX", "PDF"], default=["TXT"])
        
        st.markdown("---")
        st.info("💡 **Tip**: Upload multiple files for combined summarization!")
    
    # Main content
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📝 Input Content")
        
        input_method = st.radio("Choose input method:", ["Text Input", "File Upload"])
        
        text_content = ""
        
        if input_method == "Text Input":
            text_content = st.text_area(
                "Paste your content here:",
                height=300,
                placeholder="Enter story, article, script, research paper, or any text..."
            )
        else:
            uploaded_files = st.file_uploader(
                "Upload files (PDF, DOCX, TXT)",
                type=['pdf', 'docx', 'txt'],
                accept_multiple_files=True
            )
            
            if uploaded_files:
                combined_text = []
                for file in uploaded_files:
                    with st.spinner(f"Processing {file.name}..."):
                        if file.name.endswith('.pdf'):
                            text = extract_text_from_pdf(file)
                        elif file.name.endswith('.docx'):
                            text = extract_text_from_docx(file)
                        else:
                            text = extract_text_from_txt(file)
                        
                        if text:
                            combined_text.append(f"\n--- {file.name} ---\n{text}")
                
                text_content = "\n".join(combined_text)
                
                if text_content:
                    st.success(f"✅ Loaded {len(uploaded_files)} file(s)")
                    with st.expander("Preview extracted text"):
                        st.text(text_content[:1000] + "..." if len(text_content) > 1000 else text_content)
        
        summarize_button = st.button("🚀 Generate Summary", type="primary")
    
    with col2:
        st.subheader("✨ Results")
        
        if summarize_button and text_content:
            original_text = text_content
            processed_text = preprocess_text(text_content)
            
            if not processed_text or len(processed_text.split()) < 10:
                st.error("❌ Text too short or empty. Please provide more content.")
                return
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Step 1: Language detection and translation
            status_text.text("🔍 Detecting language...")
            progress_bar.progress(10)
            
            detected_lang = detect_language(processed_text)
            original_lang = detected_lang
            
            if enable_translation and detected_lang != 'en':
                status_text.text(f"🌍 Translating from {detected_lang} to English...")
                progress_bar.progress(20)
                processed_text = translate_text(processed_text, 'en')
            else:
                progress_bar.progress(20)
            
            # Step 2: Summarization
            status_text.text("🤖 Generating summary...")
            progress_bar.progress(40)
            
            start_time = time.time()
            
            if turbo_mode:
                summary = turbo_summarize(processed_text, num_sentences)
            else:
                summary = deep_summarize(processed_text, max_length=summary_length)
            
            time_taken = time.time() - start_time
            progress_bar.progress(60)
            
            # Step 3: Translate summary back if needed
            if enable_translation and original_lang != 'en':
                status_text.text(f"🌍 Translating summary back to {original_lang}...")
                progress_bar.progress(70)
                summary = translate_text(summary, original_lang)
            
            # Step 4: Generate insights
            status_text.text("📊 Analyzing content...")
            progress_bar.progress(80)
            
            tone = analyze_tone(processed_text)
            keywords = extract_keywords(processed_text)
            
            # Step 5: Wikipedia enrichment
            wiki_context = None
            if enable_wikipedia:
                status_text.text("📚 Fetching Wikipedia context...")
                if 'wiki_query' in locals() and wiki_query:
                    wiki_context = fetch_wikipedia_context(wiki_query)
                elif keywords:
                    wiki_context = fetch_wikipedia_context(keywords[0])
            
            progress_bar.progress(90)
            
            # Step 6: Generate Q&A notes
            status_text.text("📝 Generating study notes...")
            qa_notes = generate_qa_notes(processed_text, summary)
            
            progress_bar.progress(100)
            status_text.text("✅ Complete!")
            time.sleep(0.5)
            progress_bar.empty()
            status_text.empty()
            
            # Display results
            st.markdown("### 📄 Summary")
            st.success(summary)
            
            # Metrics
            col_m1, col_m2, col_m3 = st.columns(3)
            with col_m1:
                st.metric("Original Words", len(original_text.split()))
            with col_m2:
                st.metric("Summary Words", len(summary.split()))
            with col_m3:
                reduction = (1 - len(summary.split()) / len(original_text.split())) * 100
                st.metric("Reduction", f"{reduction:.1f}%")
            
            st.metric("⏱️ Processing Time", f"{time_taken:.2f}s")
            
            # Insights
            st.markdown("### 🔍 Insights")
            col_i1, col_i2 = st.columns(2)
            
            with col_i1:
                st.markdown(f"**Tone:** {tone}")
            
            with col_i2:
                st.markdown(f"**Top Keywords:** {', '.join(keywords[:5])}")
            
            # Wikipedia context
            if wiki_context:
                with st.expander("📚 Wikipedia Context"):
                    st.info(wiki_context)
            
            # Q&A Notes
            if qa_notes:
                with st.expander("📝 Study Notes (Q&A)"):
                    for qa in qa_notes:
                        st.markdown(f"**{qa['question']}**")
                        st.markdown(qa['answer'])
                        st.markdown("---")
            
            # Export section
            st.markdown("### 📤 Export Summary")
            
            export_content = f"""UNIVERSAL AI SUMMARIZER - REPORT
{'='*50}

SUMMARY:
{summary}

METADATA:
- Original Length: {len(original_text.split())} words
- Summary Length: {len(summary.split())} words
- Reduction: {reduction:.1f}%
- Processing Time: {time_taken:.2f}s
- Tone: {tone}
- Keywords: {', '.join(keywords)}

"""
            
            if wiki_context:
                export_content += f"\nWIKIPEDIA CONTEXT:\n{wiki_context}\n"
            
            if qa_notes:
                export_content += "\nSTUDY NOTES (Q&A):\n"
                for qa in qa_notes:
                    export_content += f"\n{qa['question']}\n{qa['answer']}\n"
            
            col_e1, col_e2, col_e3 = st.columns(3)
            
            if "TXT" in export_format:
                with col_e1:
                    txt_data = export_to_txt(export_content)
                    st.download_button(
                        "📄 Download TXT",
                        data=txt_data,
                        file_name="summary.txt",
                        mime="text/plain"
                    )
            
            if "DOCX" in export_format:
                with col_e2:
                    docx_data = export_to_docx(export_content)
                    st.download_button(
                        "📘 Download DOCX",
                        data=docx_data,
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            if "PDF" in export_format:
                with col_e3:
                    pdf_data = export_to_pdf(export_content)
                    if pdf_data:
                        st.download_button(
                            "📕 Download PDF",
                            data=pdf_data,
                            file_name="summary.pdf",
                            mime="application/pdf"
                        )
    
    # Additional features in tabs
    st.markdown("---")
    tab1, tab2, tab3 = st.tabs(["📊 Summary Evaluation", "📚 Multi-Doc Summary", "⚡ Speed Comparison"])
    
    with tab1:
        st.subheader("📊 Evaluate Summary Quality")
        st.write("Compare AI-generated summary with your own or another reference summary.")
        
        col_eval1, col_eval2 = st.columns(2)
        
        with col_eval1:
            ai_summary = st.text_area("AI Summary", height=150, placeholder="Paste AI-generated summary...")
        
        with col_eval2:
            human_summary = st.text_area("Reference Summary", height=150, placeholder="Paste reference summary...")
        
        if st.button("Evaluate Similarity"):
            if ai_summary and human_summary:
                similarity = calculate_similarity(ai_summary, human_summary)
                st.metric("Similarity Score", f"{similarity:.1f}%")
                
                if similarity > 80:
                    st.success("✅ Excellent match!")
                elif similarity > 60:
                    st.info("✓ Good match")
                else:
                    st.warning("⚠ Significant differences detected")
                
                with st.expander("View Differences"):
                    diff_html = generate_diff_html(human_summary, ai_summary)
                    st.markdown(f'<div style="padding: 1rem; background: #f0f2f6; border-radius: 5px;">{diff_html}</div>', unsafe_allow_html=True)
                    st.caption("🟢 Green = Added by AI | 🔴 Red = Missing from AI")
    
    with tab2:
        st.subheader("📚 Multi-Document Summarization")
        st.write("Upload multiple documents to generate a combined summary.")
        
        multi_files = st.file_uploader(
            "Upload multiple documents",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            key="multi_upload"
        )
        
        if multi_files and st.button("Summarize All Documents"):
            all_texts = []
            
            for file in multi_files:
                with st.spinner(f"Reading {file.name}..."):
                    if file.name.endswith('.pdf'):
                        text = extract_text_from_pdf(file)
                    elif file.name.endswith('.docx'):
                        text = extract_text_from_docx(file)
                    else:
                        text = extract_text_from_txt(file)
                    
                    if text:
                        all_texts.append(preprocess_text(text))
            
            if all_texts:
                combined = "\n\n".join(all_texts)
                
                with st.spinner("Generating combined summary..."):
                    if turbo_mode:
                        multi_summary = turbo_summarize(combined, num_sentences * len(all_texts))
                    else:
                        multi_summary = deep_summarize(combined, max_length=summary_length * 2)
                
                st.success("Combined Summary:")
                st.write(multi_summary)
                
                st.metric("Total Documents", len(all_texts))
                st.metric("Combined Length", f"{len(combined.split())} words")
    
    with tab3:
        st.subheader("⚡ Speed Comparison: Turbo vs Deep")
        st.write("Test both summarization modes on sample text to compare speed and output.")
        
        sample_text = st.text_area(
            "Sample Text for Speed Test",
            height=150,
            value="Artificial intelligence is revolutionizing the way we interact with technology. Machine learning algorithms can now process vast amounts of data to identify patterns and make predictions. Natural language processing enables computers to understand and generate human language. Computer vision allows machines to interpret and analyze visual information from the world. These technologies are being applied in healthcare, finance, transportation, and many other fields to solve complex problems and improve efficiency."
        )
        
        if st.button("Run Speed Test") and sample_text:
            col_speed1, col_speed2 = st.columns(2)
            
            with col_speed1:
                st.markdown("**🚀 Turbo Mode (LSA)**")
                start = time.time()
                turbo_result = turbo_summarize(sample_text, 3)
                turbo_time = time.time() - start
                st.write(turbo_result)
                st.metric("Time", f"{turbo_time:.3f}s")
            
            with col_speed2:
                st.markdown("**🧠 Deep Mode (Transformer)**")
                start = time.time()
                deep_result = deep_summarize(sample_text, max_length=100)
                deep_time = time.time() - start
                st.write(deep_result)
                st.metric("Time", f"{deep_time:.3f}s")
            
            speedup = deep_time / turbo_time
            st.info(f"⚡ Turbo mode is **{speedup:.1f}x faster** than Deep mode!")
    
    # Footer
    st.markdown("---")
    st.markdown(
        '<div class="footer">🧠 Built by Aditya — Universal AI Summarizer v4.1 Turbo | Hugging Face + Streamlit + Free APIs</div>',
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
